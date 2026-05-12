"""DOCX parser using python-docx. Returns paragraphs and table cells."""
from docx import Document


def parse_docx(path: str) -> tuple[list[str], bool]:
    """Returns (paragraphs: list[str], ocr_used=False)."""
    doc = Document(path)
    texts: list[str] = []

    # Body paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            texts.append(text)

    # Table cells (useful for structured content like syllabi)
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                texts.append(" | ".join(row_texts))

    return texts, False
